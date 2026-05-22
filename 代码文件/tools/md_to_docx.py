#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用 Markdown → DOCX 转换器
使用 python-docx 库，支持标题/段落/粗体/表格/列表/水平线/代码块/引用块。

用法:
    python md_to_docx.py <input.md> <output.docx> [title]

依赖:
    pip install python-docx
"""

import re
import sys
import os

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 颜色常量 ──────────────────────────────────────────────
C_PRIMARY     = RGBColor(0x1A, 0x1A, 0x2E)  # h1 颜色 / 表头背景
C_SECONDARY   = RGBColor(0x16, 0x21, 0x3E)  # h2 颜色
C_BODY        = RGBColor(0x33, 0x33, 0x33)  # h3 / 正文颜色
C_GRAY        = RGBColor(0x55, 0x55, 0x55)  # 引用块颜色
C_CODE_BG     = RGBColor(0xE7, 0x4C, 0x3C)  # 行内代码文字色
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

FONT_FAMILY   = 'Microsoft YaHei'
FONT_MONO     = 'Consolas'
SIZE_NORMAL   = Pt(10.5)
SIZE_H1       = Pt(18)
SIZE_H2       = Pt(14)
SIZE_H3       = Pt(12)
SIZE_SMALL    = Pt(9)


# ── 辅助函数 ──────────────────────────────────────────────

def _set_east_asian_font(rPr, font_name):
    """Set w:eastAsia font on an rPr element."""
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def set_run_font(run, font_name=FONT_FAMILY, size=SIZE_NORMAL,
                  color=None, bold=False):
    """Configure font properties on a single run."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Ensure east-asia font is set
    rPr = run._element.get_or_add_rPr()
    _set_east_asian_font(rPr, font_name)


def configure_style(style, font_name=FONT_FAMILY, size=SIZE_NORMAL,
                    color=None, bold=False):
    """Configure font properties on a named style."""
    style.font.name = font_name
    style.font.size = size
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    # East-asia font on the style's rPr
    rPr = style.element.rPr
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{font_name}"/></w:rPr>')
        style.element.append(rPr)
    else:
        _set_east_asian_font(rPr, font_name)


def add_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def split_formatted_text(text):
    """Split text into segments: (plain_string, is_bold, is_code).

    Returns list of dicts: {'text': str, 'bold': bool, 'code': bool}
    """
    segments = []
    # Pattern: **bold** or `inline code`
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            segments.append({'text': part[2:-2], 'bold': True, 'code': False})
        elif part.startswith('`') and part.endswith('`'):
            segments.append({'text': part[1:-1], 'bold': False, 'code': True})
        else:
            segments.append({'text': part, 'bold': False, 'code': False})
    return segments


def add_formatted_paragraph(doc, text, style_name='Normal',
                            font_name=FONT_FAMILY, size=SIZE_NORMAL,
                            color=None, bold=False, alignment=None,
                            left_indent=None, italic=False):
    """Add a paragraph with bold/inline-code formatting support."""
    p = doc.add_paragraph(style=style_name)
    if alignment:
        p.alignment = alignment
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    p.paragraph_format.space_after = Pt(4)

    segments = split_formatted_text(text)
    for seg in segments:
        run = p.add_run(seg['text'])
        seg_bold = bold or seg['bold']
        if seg['code']:
            set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL,
                         color=C_CODE_BG, bold=seg_bold)
        else:
            set_run_font(run, font_name=font_name, size=size,
                         color=color, bold=seg_bold)
        if italic or (style_name == 'Normal' and seg == segments[0] and color == C_GRAY):
            pass  # italic handled below
        run.font.italic = italic

    return p


def _is_separator_row(stripped):
    """Check if a table line is a separator (e.g. |---|---|)."""
    if not (stripped.startswith('|') and stripped.endswith('|')):
        return False
    inner = stripped[1:-1].replace(' ', '')
    return all(ch in '-:|' for ch in inner) and '-' in inner


def _parse_table_cells(line):
    """Split a markdown table row into cell strings."""
    cells = [c.strip() for c in line.strip().split('|')]
    # Remove leading/trailing empty cells from split at start/end
    if cells and cells[0] == '':
        cells = cells[1:]
    if cells and cells[-1] == '':
        cells = cells[:-1]
    return cells


def render_table(doc, raw_lines):
    """Convert a batch of consecutive markdown table lines to a Word table.

    raw_lines: list of strings like '| a | b | c |'
    """
    # Filter out separator rows (e.g. |---|---|)
    data = [ln for ln in raw_lines if not _is_separator_row(ln)]
    if len(data) < 2:
        return  # need at least header + one data row

    rows_data = [_parse_table_cells(ln) for ln in data]
    num_cols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j in range(num_cols):
            cell = row.cells[j]
            cell.text = ''
            if j < len(row_data):
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                segments = split_formatted_text(row_data[j])
                for seg in segments:
                    run = p.add_run(seg['text'])
                    if seg['code']:
                        set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL,
                                     color=C_CODE_BG)
                    else:
                        set_run_font(run, bold=seg['bold'])

            # Header row: dark background, white bold text
            if i == 0:
                add_cell_shading(cell, '1A1A2E')
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = C_WHITE
                    run.font.bold = True

    # Small spacer after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 主转换函数 ────────────────────────────────────────────

def md_to_docx(md_path: str, docx_path: str, title: str = None) -> None:
    """Convert a Markdown file to a styled DOCX document."""

    # Read markdown — use utf-8-sig to auto-strip BOM if present
    with open(md_path, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()

    doc = Document()

    # ── Configure built-in styles ──────────────────────────
    configure_style(doc.styles['Normal'],
                    font_name=FONT_FAMILY, size=SIZE_NORMAL)
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.15

    heading_config = {
        1: (SIZE_H1, C_PRIMARY),
        2: (SIZE_H2, C_SECONDARY),
        3: (SIZE_H3, C_BODY),
    }
    for level, (size, color) in heading_config.items():
        configure_style(doc.styles[f'Heading {level}'],
                        font_name=FONT_FAMILY, size=size,
                        color=color, bold=True)

    # ── Optional document title ───────────────────────────
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, font_name=FONT_FAMILY, size=SIZE_H1,
                     color=C_PRIMARY, bold=True)
        p.paragraph_format.space_after = Pt(12)

    # ── State ─────────────────────────────────────────────
    in_code_block = False
    code_lines = []

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip('\n')
        s = raw.strip()

        # ── Code block ────────────────────────────────────
        if s.startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────
        if not s:
            i += 1
            continue

        # ── Thematic break (---) ──────────────────────────
        if re.match(r'^-{3,}\s*$', s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────
        if s.startswith('> '):
            q_lines = [s[2:]]
            i += 1
            while i < len(lines):
                ns = lines[i].strip()
                if ns.startswith('> '):
                    q_lines.append(ns[2:])
                    i += 1
                elif ns == '>':
                    q_lines.append('')
                    i += 1
                else:
                    break
            combined = '\n'.join(q_lines)
            p = add_formatted_paragraph(doc, combined)
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = C_GRAY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            continue

        # ── Table ─────────────────────────────────────────
        if s.startswith('|') and s.count('|') >= 2:
            rows = [s]
            i += 1
            while i < len(lines):
                ns = lines[i].strip()
                if ns.startswith('|') and ns.count('|') >= 2:
                    rows.append(ns)
                    i += 1
                else:
                    break
            render_table(doc, rows)
            continue

        # ── Headings ──────────────────────────────────────
        h_match = re.match(r'^(#{1,3})\s+(.+)$', s)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Unordered list ────────────────────────────────
        ul_match = re.match(r'^[-*]\s+(.+)$', s)
        if ul_match:
            text = ul_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            segments = split_formatted_text(text)
            for seg in segments:
                run = p.add_run(seg['text'])
                if seg['code']:
                    set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL,
                                 color=C_CODE_BG)
                else:
                    set_run_font(run, bold=seg['bold'])
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', s)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            segments = split_formatted_text(text)
            for seg in segments:
                run = p.add_run(seg['text'])
                if seg['code']:
                    set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL,
                                 color=C_CODE_BG)
                else:
                    set_run_font(run, bold=seg['bold'])
            i += 1
            continue

        # ── Plain paragraph ───────────────────────────────
        add_formatted_paragraph(doc, raw)
        i += 1

    # Flush remaining code block (no closing ```)
    if in_code_block and code_lines:
        text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, font_name=FONT_MONO, size=SIZE_SMALL)

    # ── Save ──────────────────────────────────────────────
    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    doc.save(docx_path)


# ── CLI 入口 ──────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(f'Usage: python {sys.argv[0]} <input.md> <output.docx> [title]')
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    doc_title = sys.argv[3] if len(sys.argv) > 3 else None

    if not os.path.isfile(md_path):
        print(f'[错误] 未找到文件: {md_path}')
        sys.exit(1)

    md_to_docx(md_path, docx_path, doc_title)
    print(f'[OK] {docx_path} 已生成')
