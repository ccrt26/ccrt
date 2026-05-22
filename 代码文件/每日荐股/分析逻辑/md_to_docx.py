#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 .md 转为 .docx"""
import re, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def md_to_docx(md_path, docx_path):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    lines = text.split('\n')
    n = len(lines)
    idx = 0

    while idx < n:
        line = lines[idx].rstrip()

        # Skip frontmatter
        if line.startswith('---') and idx == 0:
            idx += 1
            while idx < n and not lines[idx].strip().startswith('---'):
                idx += 1
            idx += 1
            continue

        # Skip empty lines
        if not line.strip():
            idx += 1
            continue

        # Table detection
        if line.startswith('|') and line.endswith('|'):
            # Gather all table rows
            table_rows = []
            while idx < n and lines[idx].strip().startswith('|'):
                row = lines[idx].strip()
                # Skip separator rows like |---|---|
                if not re.match(r'^[\s\|:\-]+$', row):
                    cols = [c.strip() for c in row.split('|')[1:-1]]
                    table_rows.append(cols)
                idx += 1

            if len(table_rows) >= 2:
                headers = table_rows[0]
                data = table_rows[1:]
                num_cols = max(len(headers), max(len(r) for r in data) if data else 0)

                t = doc.add_table(rows=1 + len(data), cols=num_cols)
                t.style = 'Table Grid'
                t.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Headers
                for j, h in enumerate(headers):
                    cell = t.rows[0].cells[j]
                    cell.text = h
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, '1A1A2E')

                # Data rows
                for ri, row in enumerate(data):
                    for j, val in enumerate(row):
                        if j < num_cols:
                            cell = t.rows[ri + 1].cells[j]
                            cell.text = val
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in p.runs:
                                    run.font.size = Pt(9)

                doc.add_paragraph()
            continue

        # Headings
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            if level <= 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            idx += 1
            continue

        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            idx += 1
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            idx += 1
            continue

        # Regular paragraph (collect consecutive lines)
        para_lines = [line]
        idx += 1
        while idx < n:
            next_line = lines[idx].rstrip()
            if next_line.startswith('#') or next_line.startswith('|') or not next_line.strip():
                break
            if next_line.startswith('> ') or next_line.startswith('- ') or next_line.startswith('* '):
                break
            para_lines.append(next_line)
            idx += 1

        para_text = '\n'.join(para_lines)
        # Skip separator lines
        if para_text.strip().startswith('---') or para_text.strip().startswith('==='):
            continue
        doc.add_paragraph(para_text)

    doc.save(docx_path)
    sz = os.path.getsize(docx_path)
    print(f"DOCX generated: {docx_path} ({sz} bytes)")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <md_path> <docx_path>")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
